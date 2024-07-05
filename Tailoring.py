import openai
import pandas as pd
import docx
from docx import Document
import re
import json




def extract_text(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


# CV GENERATION

def update_cv_sections(cv_json, job_description):
    prompt_text = (
        f"Given the job description below, update the json CV activities accordingly. "
        f"Ensure you produce exactly 3 bullet points for the experience section "
        f"and 2 bullet points for the projects section. "
        f"Maintain the main ideas, writing style and lengths (LESS THAN 16 WORDS) of the original activities. Most importantly use the ACTION VERBS "
        f"AND KEYWORDS FROM THE JOB OFFER when possible. Do not make drastic changes and do not repeat the same action verb more than twice. OUTPUT A VALID JSON WITH THE SAME INPUT FORMAT WITH THE UPDATED INFORMATION\n\n"
        f"Job Description: {job_description}\n\n"
        f"CV Experience Section: {json.dumps(cv_json)}")

    # API call
    response = openai.chat.completions.create(
        model='gpt-4o',
        messages=[{"role": "system",
                   "content": "You are a recruiter with 20 years of experience in big tech companies, expert in CV tailoring."},
                  {"role": "user", "content": prompt_text}],
        max_tokens=600)

    new_bullet_points = response.choices[0].message.content

    if new_bullet_points.startswith("```json"):
        new_bullet_points = new_bullet_points[7:]
    if new_bullet_points.endswith("```"):
        new_bullet_points = new_bullet_points[:-3]

    updated_cv_json = json.loads(new_bullet_points)
    # print(new_bullet_points)  # debug line
    updated_cv_json = json.loads(new_bullet_points)
    return updated_cv_json


def extract_activities(new_bullets):
    experience_activities = []
    relevant_projects_activities = []

    # Extract activities from experience
    for job in new_bullets["experience"]:
        for activity in job["activities"]:
            experience_activities.append(activity)

    # Extract activities from relevant projects
    for project in new_bullets["relevant_projects"]:
        for activity in project["activities"]:
            relevant_projects_activities.append(activity)

    return {
        "experience": experience_activities,
        "relevant_projects": relevant_projects_activities}


def update_cv(cv_path, new_bullets):
    doc = Document(cv_path)
    capture_experience = False
    capture_projects = False
    exp_index = 0
    proj_index = 0

    def count_words(line):
        words = re.findall(r'\b\w+\b', line)
        return len(words)

    # iterate through paragraphs in the doc
    for para in doc.paragraphs:
        text = para.text.strip()

        # capturing lines after "WORK EXPERIENCE"
        if text.lower() == 'work experience':
            capture_experience = True
            capture_projects = False
            continue

        # capturing lines after "RELEVANT PROJECTS"
        if text.lower() == 'relevant projects':
            capture_experience = False
            capture_projects = True
            continue

        # capture and write (preserving format) for experience
        if capture_experience and count_words(text) >= 11 and '\t' not in text:
            if exp_index < len(new_bullets['experience']):
                # get format
                if para.runs:
                    original_run = para.runs[0]
                    font_name = original_run.font.name
                    font_size = original_run.font.size
                    bold = original_run.bold
                    italic = original_run.italic

                # write a bullet point
                new_text = '• ' + new_bullets['experience'][exp_index]
                p = para._element
                for child in p[:]:
                    p.remove(child)
                run = para.add_run(new_text)

                # formatting
                if para.runs:
                    run.font.name = font_name
                    run.font.size = font_size
                    run.bold = bold
                    run.italic = italic

                exp_index += 1

        # capture and write (preserving format) for relevant projects
        if capture_projects and count_words(text) >= 11 and '\t' not in text:
            if proj_index < len(new_bullets['relevant_projects']):
                # get format
                if para.runs:
                    original_run = para.runs[0]
                    font_name = original_run.font.name
                    font_size = original_run.font.size
                    bold = original_run.bold
                    italic = original_run.italic

                # write a bullet point
                new_text = '• ' + new_bullets['relevant_projects'][proj_index]
                p = para._element
                for child in p[:]:
                    p.remove(child)
                run = para.add_run(new_text)

                # formatting
                if para.runs:
                    run.font.name = font_name
                    run.font.size = font_size
                    run.bold = bold
                    run.italic = italic

                proj_index += 1

    # doc.save('updated_document_with_formatting.docx')
    return doc


def generate_cv(cv_path, cv_json, job_offer):
    updated_cv_json = update_cv_sections(cv_json, job_offer)
    new_bullets = extract_activities(updated_cv_json)
    return update_cv(cv_path, new_bullets)


# COVER LETTER GENERATION
def generate_cover_letter(company, position, job_offer):
    cover_letter_path = 'Cover Letter.docx'
    # cover letter text
    cover_letter_template = extract_text(cover_letter_path)

    # prompt for GPT-4o
    prompt_text = (
        f"Using the cover letter template and Job Offer Details provided below, craft a complete, compelling cover letter. "
        f"The cover letter should articulate why the applicant is drawn to {company}, and fit to be a {position}and showcase my skills, experiences and competences "
        f"Leverage all my qualities, also showing my softskills, and dont invent false information "
        f"specifically citing relevant company values and initiatives mentioned in the job offer but mainly ones found online. "
        f"Keep it under 470 words so trim information on the template that is not relevant to the job offer and make sure you return a cover letter from a person you would definitely hire."
        f"Maintain a professional writing style and match the writing style of the template"
        f"Moreover, I am attacching bullet points from my cv, please adapt them with the info of the cover letter and my cover letter\n\n"
        f"Job Offer Details: {job_offer}\n\n"
        f"Cover Letter Template: {cover_letter_template}")

    # API call
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system",
                   "content": f"You are a recruiter with 20 years of experience in big tech companies, expert in CV and cover letter taloring"},
                  {"role": "user", "content": prompt_text}],
        max_tokens=620)

    return response.choices[0].message.content


def tailor():
    job_df = pd.read_csv('Cover Letter List.csv', encoding='latin-1')
    for index, row in job_df.iterrows():
        company = row['Company']
        position = row['Position']
        job_offer = row['Job Offer']
        needs_cover_letter = row['Needs Cover Letter']

        # CV
        cv_path = 'CV Rodrigo Ugarte.docx'
        cv_json = json.load(open('cv_data.json', encoding='utf-8'))
        cv = generate_cv(cv_path, cv_json, job_offer)
        cv.save(f'CV_Rodrigo_Ugarte_{company}_{position}.docx')
        print(f'CV for {company} for {position} done!')

        # COVER LETTER
        if needs_cover_letter:
            cover_letter = generate_cover_letter(company, position, job_offer)
            doc = Document()
            doc.add_paragraph(cover_letter)
            doc.save(f'Cover_Letter_{company}_{position}.docx')
            print(f'Cover Letter for {company} for {position} done!')

tailor()